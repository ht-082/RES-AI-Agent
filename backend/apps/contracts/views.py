from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from .models import ContractTemplate, ContractDraft, ContractReview, ContractReviewFinding
from .serializers import (
    ContractTemplateSerializer, ContractTemplateDetailSerializer,
    ContractDraftSerializer, ContractReviewSerializer,
)


class ContractTemplateViewSet(viewsets.ReadOnlyModelViewSet):
    """표준 계약서 종류 조회 (읽기 전용)"""
    queryset = ContractTemplate.objects.filter(is_active=True)
    serializer_class = ContractTemplateSerializer
    lookup_field = 'code'
    filterset_fields = ['category']

    def get_serializer_class(self):
        if self.action == 'retrieve':
            return ContractTemplateDetailSerializer
        return ContractTemplateSerializer


class ContractDraftViewSet(viewsets.ModelViewSet):
    """계약서 신규 생성 (K-1)"""
    queryset = ContractDraft.objects.all()
    serializer_class = ContractDraftSerializer

    def create(self, request, *args, **kwargs):
        """POST /api/contracts/drafts/ — 계약서 초안 생성"""
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        draft = serializer.save(status='generating')

        # Mock 생성 (추후 LLM 호출로 교체)
        template = draft.template
        key_terms = draft.key_terms

        mock_content = self._generate_mock_draft(template, key_terms)
        draft.generated_content = mock_content
        draft.status = 'completed'
        draft.save()

        return Response(
            ContractDraftSerializer(draft).data,
            status=status.HTTP_201_CREATED
        )

    @action(detail=True, methods=['get'], url_path='download')
    def download(self, request, pk=None):
        """GET /api/contracts/drafts/{id}/download — Word 다운로드"""
        draft = self.get_object()
        # TODO: Word 파일 생성 로직 (python-docx)
        return Response({
            'message': '다운로드 기능은 구현 예정입니다.',
            'draft_id': str(draft.id),
        })

    def _generate_mock_draft(self, template, key_terms):
        """Mock 계약서 초안 생성"""
        parties = key_terms.get('parties', '(주)재생E파워 / ○○에너지(주)')
        capacity = key_terms.get('capacity', '80 MW')
        period = key_terms.get('period', '20년')

        return f"""# {template.name_ko}

## 계약 당사자
{parties}

## 제1조 (목적)
본 계약은 {capacity} 규모의 재생에너지 발전설비에 관한 권리·의무를 정합니다.

## 제2조 (계약 기간)
상업운전 개시일로부터 {period}으로 합니다.

## 제3조 (정산 조건)
{key_terms.get('pricing', '별도 협의')}

---
※ 본 초안은 시스템 개발 중 생성된 Mock 데이터입니다.
   LLM API 연결 후 실제 계약서 양식 기반 초안이 생성됩니다.
"""


class ContractReviewViewSet(viewsets.ModelViewSet):
    """계약서 검토 (K-2)"""
    queryset = ContractReview.objects.all()
    serializer_class = ContractReviewSerializer

    def create(self, request, *args, **kwargs):
        """POST /api/contracts/reviews/ — 계약서 검토 실행"""
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)

        review = serializer.save(status='reviewing')

        # Mock 검토 결과 생성 (추후 LLM 호출로 교체)
        mock_findings = self._generate_mock_findings()
        for finding_data in mock_findings:
            ContractReviewFinding.objects.create(review=review, **finding_data)

        review.summary = '총 3건의 지적사항이 발견되었습니다. 독소조항 1건, 불리조항 1건, 누락 1건.'
        review.status = 'completed'
        review.save()

        return Response(
            ContractReviewSerializer(review).data,
            status=status.HTTP_201_CREATED
        )

    @action(detail=True, methods=['get'], url_path='download')
    def download(self, request, pk=None):
        """GET /api/contracts/reviews/{id}/download — Word 다운로드"""
        review = self.get_object()
        return Response({
            'message': '다운로드 기능은 구현 예정입니다.',
            'review_id': str(review.id),
        })

    def _generate_mock_findings(self):
        """Mock 검토 결과 데이터"""
        return [
            {
                'clause_ref': '제12조',
                'severity': 'high',
                'category': '독소조항',
                'finding': '지체상금 상한 없음',
                'suggestion': '계약금액의 10% 상한 신설 제안',
                'order_index': 1,
            },
            {
                'clause_ref': '제18조',
                'severity': 'mid',
                'category': '불리조항',
                'finding': '하자담보 책임 5년 → 표준 2년 대비 과도',
                'suggestion': '단축 협상 필요',
                'order_index': 2,
            },
            {
                'clause_ref': '—',
                'severity': 'low',
                'category': '누락',
                'finding': '불가항력 조항 부재',
                'suggestion': '표준 양식 제24조 삽입 권장',
                'order_index': 3,
            },
        ]


import os
import json
from django.conf import settings
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated

class ContractTypesAPIView(APIView):
    """GET /api/contract/types — 등록된 계약 유형 목록과 fields 스키마 반환"""
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        json_path = os.path.join(settings.BASE_DIR, 'apps', 'contracts', 'contract_types.json')
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return Response(data, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(
                {"error": f"계약 유형 설정을 읽어오지 못했습니다. {e}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class ContractGenerateAPIView(APIView):
    """POST /api/contract/generate — 입력 조건에 따른 조항 기반 초안 생성"""
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        type_id = request.data.get('type_id')
        inputs = request.data.get('inputs', {})

        if not type_id:
            return Response({"error": "type_id는 필수 항목입니다."}, status=status.HTTP_400_BAD_REQUEST)

        # JSON 설정에서 type_id 찾기
        json_path = os.path.join(settings.BASE_DIR, 'apps', 'contracts', 'contract_types.json')
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                types_data = json.load(f)
        except Exception as e:
            return Response({"error": f"계약 유형 설정을 읽지 못했습니다. {e}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        selected_type = next((t for t in types_data if t['type_id'] == type_id), None)
        if not selected_type:
            return Response({"error": f"존재하지 않는 계약 유형입니다: {type_id}"}, status=status.HTTP_400_BAD_REQUEST)

        # 필수 필드 검증
        for field in selected_type.get('fields', []):
            if field.get('required') and not inputs.get(field['key']):
                return Response(
                    {"error": f"필수 입력 항목 누락: {field['label']}({field['key']})"},
                    status=status.HTTP_400_BAD_REQUEST
                )

        # LLM 호출
        from services.llm import generate_structured_contract
        
        result_json = generate_structured_contract(
            type_id=type_id,
            inputs=inputs,
            article_structure=selected_type['article_structure'],
            system_prompt=selected_type['system_prompt']
        )

        # DB에 생성 이력 저장
        draft = ContractDraft.objects.create(
            user=request.user,
            type_id=type_id,
            inputs=inputs,
            generated_articles=result_json.get('articles', []),
            title=result_json.get('title', f"{selected_type['type_name']} 초안"),
            status='completed'
        )

        # 최종 응답 데이터 구성
        response_data = {
            "draft_id": str(draft.id),
            "title": result_json.get('title'),
            "articles": result_json.get('articles', []),
            "mapping_note": result_json.get('mapping_note', '')
        }

        return Response(response_data, status=status.HTTP_201_CREATED)


from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from urllib.parse import quote
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rest_framework.exceptions import PermissionDenied

class ContractDraftDownloadAPIView(APIView):
    """GET /api/contract/drafts/<draft_id>/download — 생성된 초안을 Word(.docx) 파일로 다운로드"""
    permission_classes = [IsAuthenticated]

    def get(self, request, draft_id, *args, **kwargs):
        draft = get_object_or_404(ContractDraft, id=draft_id)

        # 본인이 생성한 draft만 다운로드 허용
        if draft.user != request.user:
            raise PermissionDenied("본인이 생성한 계약서 초안만 다운로드할 수 있습니다.")

        # python-docx를 사용하여 문서 작성
        doc = Document()

        # 1. 문서 제목 (가운데 정렬, 큰 글씨)
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(draft.title or "계약서 초안")
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_p.paragraph_format.space_after = Pt(24)

        # 2. 조항별 heading + 본문 문단
        articles = draft.generated_articles or []
        for art in articles:
            no = art.get('no', '')
            heading = art.get('heading')
            content = art.get('content', '')

            # Heading (제N조 + 제목) 스타일
            heading_text = f"{no}"
            if heading:
                heading_text += f" ({heading})"

            heading_p = doc.add_paragraph()
            heading_run = heading_p.add_run(heading_text)
            heading_run.font.bold = True
            heading_run.font.size = Pt(12)
            heading_p.paragraph_format.space_before = Pt(14)
            heading_p.paragraph_format.space_after = Pt(6)

            # 본문 문단
            content_p = doc.add_paragraph()
            content_run = content_p.add_run(content)
            content_run.font.size = Pt(10)
            content_p.paragraph_format.left_indent = Inches(0.2)
            content_p.paragraph_format.space_after = Pt(12)

        # 3. 마지막 면책 문구 문단
        doc.add_paragraph()  # 여백 한 줄
        disclaimer_p = doc.add_paragraph()
        disclaimer_run = disclaimer_p.add_run("※ 본 초안은 참고용이며, 최종 계약은 법무 검토가 필요합니다.")
        disclaimer_run.font.size = Pt(9.5)
        disclaimer_run.font.italic = True
        disclaimer_p.paragraph_format.space_before = Pt(24)

        # 메모리 파일로 저장
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # 파일명 구성: "{계약유형}_초안_{날짜}.docx"
        type_mapping = {
            'ppa': '전력판매계약(PPA)',
            'epc': 'EPC도급계약',
            'om': 'O&M위탁계약',
            'spa': '사업권양수도계약',
            'lease': '토지임대차계약'
        }
        type_ko = type_mapping.get(draft.type_id.lower(), draft.type_id.upper())
        date_str = datetime.now().strftime("%Y%m%d")
        filename = f"{type_ko}_초안_{date_str}.docx"

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f"attachment; filename*=utf-8''{quote(filename)}"
        return response
